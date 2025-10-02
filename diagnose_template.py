from docx import Document
import re
import os

def diagnose_placeholders(doc_path, replacements):
    """診斷 Word 文檔中的佔位符"""
    
    # 檢查檔案是否存在
    if not os.path.exists(doc_path):
        print(f"錯誤: 檔案 {doc_path} 不存在")
        return
    
    try:
        doc = Document(doc_path)
        found_placeholders = set()
        missing_placeholders = set()
        
        # 正則表達式來匹配 {{xxx}} 格式的佔位符
        placeholder_pattern = r'\{\{(.+?)\}\}'
        
        # 檢查段落中的佔位符
        for paragraph in doc.paragraphs:
            for match in re.finditer(placeholder_pattern, paragraph.text):
                placeholder = match.group(0)
                if placeholder in replacements:
                    found_placeholders.add(placeholder)
                else:
                    missing_placeholders.add(placeholder)
        
        # 檢查表格中的佔位符
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for paragraph in cell.paragraphs:
                        for match in re.finditer(placeholder_pattern, paragraph.text):
                            placeholder = match.group(0)
                            if placeholder in replacements:
                                found_placeholders.add(placeholder)
                            else:
                                missing_placeholders.add(placeholder)
        
        # 檢查替換字典中定義但在文檔中找不到的佔位符
        unused_replacements = set(replacements.keys()) - found_placeholders
        
        # 檢查格式不匹配的佔位符（大小寫問題）
        mismatched_placeholders = {}
        for unused in unused_replacements:
            # 尋找可能的大小寫不匹配
            all_doc_placeholders = found_placeholders.union(missing_placeholders)
            similar_placeholders = [p for p in all_doc_placeholders 
                                  if p.lower() == unused.lower() and p != unused]
            if similar_placeholders:
                mismatched_placeholders[unused] = similar_placeholders
        
        # 輸出診斷結果
        print(f"診斷結果 for {doc_path}:")
        print("=" * 50)
        print(f"✅ 成功匹配的佔位符 ({len(found_placeholders)} 個):")
        for placeholder in sorted(found_placeholders):
            print(f"   {placeholder}")
        
        print(f"\n⚠️  文檔中存在但未在替換字典中定義的佔位符 ({len(missing_placeholders)} 個):")
        for placeholder in sorted(missing_placeholders):
            print(f"   {placeholder}")
        
        print(f"\n❌ 替換字典中定義但在文檔中找不到的佔位符 ({len(unused_replacements)} 個):")
        for placeholder in sorted(unused_replacements):
            print(f"   {placeholder}")
        
        if mismatched_placeholders:
            print(f"\n🔍 可能存在大小寫或格式問題的佔位符:")
            for expected, found_list in mismatched_placeholders.items():
                print(f"   預期: {expected}")
                print(f"   文檔中發現: {found_list}")
        
        print("\n" + "=" * 50)
        
    except Exception as e:
        print(f"錯誤: 無法處理檔案 {doc_path}")
        print(f"錯誤詳情: {str(e)}")

def main():
    """主函數"""
    replacements = {
        '{{申請單編號}}': 'DCR-20250812-01',
        '{{申請單位}}': '信用卡行企',
        '{{申請日期}}': '2025-08-12',
        '{{申請單位填表人}}': '曹世明',
        '{{變更類型}}': 'DCR',
        '{{程式負責人}}': '優利/Daniel Chou',
        '{{預計完成日}}': '2025/09/09',
        '{{測試日期}}': '2025/09/09',
        '{{上線日}}': '2025/10/03',
        '{{進館通知函日期}}': '2025年10月03日',
        '{{內容說明}}': """
1.	CTI→受理詐騙帳戶交易→帳戶設定→警示帳戶凍結
原使用164C，調整改用16W9（新增一欄 理由碼）
2.	CTI→受理詐騙帳戶交易→帳戶設定→自轉戶設定（警示帳戶19）
原使用164C，調整改用16W9（新增一欄 理由碼）
3.	CTI→受理詐騙帳戶交易→帳戶設定→自轉戶解除（風險帳戶29）
原使用165C，調整改用16W9（新增一欄 理由碼）
4.	對應的MW報表查詢1-3、1-4的修改。


""",
        '{{上線清單}}': """
[前端js]
•	Web/vue/dist/bankAgent.1.0.146.js
•	Web/vue/dist/cardAgent.1.0.146.js
•	Web/vue/dist/collectAgent.1.0.146.js
•	Web/vue/dist/supervisor.1.0.146.js
•	Web/vue/dist/vendors.1.0.146.js

[Web]
•	Web/bin/Web.dll

[Report]
•	Web/Report/bkmw-a4.rdlc

""",
        '{{過版流程}}': """

Web前端JS異動：
1.	複製web目錄下vue目錄，進入172.30.101.11 主機，開啟檔案總管，進入d:\\web\\目錄之下後，直接覆蓋vue目錄。
2.	複製web目錄下vue目錄，進入172.30.101.21 主機，開啟檔案總管，進入d:\\web\\目錄之下後，直接覆蓋vue目錄。
3.	複製web目錄下vue目錄，進入172.30.101.31 主機，開啟檔案總管，進入d:\\web\\目錄之下後，直接覆蓋vue目錄。
4.	複製web目錄下vue目錄，進入172.30.101.41 主機，開啟檔案總管，進入d:\\web\\目錄之下後，直接覆蓋vue目錄。

Web專案dll異動：
1.	複製web目錄下bin目錄，進入172.30.101.11 主機，開啟檔案總管，進入d:\\web\\bin目錄之下後，直接覆蓋bin目錄。
2.	複製web目錄下bin目錄，進入172.30.101.21 主機，開啟檔案總管，進入d:\\web\\bin目錄之下後，直接覆蓋bin目錄。
3.	複製web目錄下bin目錄，進入172.30.101.31 主機，開啟檔案總管，進入d:\\web\\bin目錄之下後，直接覆蓋bin目錄。
4.	複製web目錄下bin目錄，進入172.30.101.41 主機，開啟檔案總管，進入d:\\web\\bin目錄之下後，直接覆蓋bin目錄。
5.	以上動作完成。

SQL執行：
使用SSMS連線至172.30.101.62 > DB[ ucti ]，開啟檔案全選執行=>
•	[聯邦行企]240627-行企查不到NEW客戶-NEW部分欄位不能為null.sql

報表rdlc替換：
4.	複製web目錄下report目錄，進入172.30.101.41 主機，開啟檔案總管，進入d:\\web\\report 目錄之下，將上述要異動的*.rdlc覆蓋。
""",
        
        '{{測試項目}}': """
修改整體CTI資料庫查詢的效能優化後的問題：
1.	無法查詢NEW客戶資料。
2.  其他。
""",
        '{{處理方法說明}}': """
新增功能：
1.針對MW有關於16W9修改對應的後端、前端功能。
2.針對MW有關於16W9修改對應的報表功能。

""",
    }

    # 檢查 template 目錄是否存在
    template_dir = 'template'
    if not os.path.exists(template_dir):
        print(f"錯誤: {template_dir} 目錄不存在")
        return

    templates = [
        'template/設計變更申請單.docx',
        'template/系統過版申請單.docx',
        'template/系統測試表.docx',
        'template/聯邦銀行廠商進館上線函.docx'
    ]

    print("開始診斷 Word 模板中的佔位符...")
    print("=" * 60)
    
    for template in templates:
        diagnose_placeholders(template, replacements)
        print()

if __name__ == '__main__':
    main()
