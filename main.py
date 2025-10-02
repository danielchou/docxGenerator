from docx import Document
import os
import yaml
import glob

def replace_placeholders(doc, replacements):
    """替換文檔中的佔位符"""
    for paragraph in doc.paragraphs:
        for run in paragraph.runs:
            for key, value in replacements.items():
                if key in run.text:
                    run.text = run.text.replace(key, str(value))
    
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        for key, value in replacements.items():
                            if key in run.text:
                                run.text = run.text.replace(key, str(value))

def process_template(template_path, base_output_dir, application_data):
    """處理單個模板檔案"""
    try:
        doc = Document(template_path)
        replace_placeholders(doc, application_data['replacements'])
        
        application_number = application_data['申請單編號']
        application_unit = application_data['申請單位']
        output_dir = os.path.join(base_output_dir, application_number)
        os.makedirs(output_dir, exist_ok=True)
        
        template_name = os.path.splitext(os.path.basename(template_path))[0]
        output_filename = f"{application_unit}-{template_name}-{application_number}.docx"
        output_path = os.path.join(output_dir, output_filename)
        
        doc.save(output_path)
        print(f"處理完成: {output_path}")
        
    except Exception as e:
        print(f"錯誤: 無法處理模板 {template_path}")
        print(f"錯誤詳情: {str(e)}")

def load_application_data(yaml_file):
    """載入申請單資料"""
    try:
        with open(yaml_file, 'r', encoding='utf-8') as f:
            return yaml.safe_load(f)
    except Exception as e:
        print(f"錯誤: 無法讀取 YAML 檔案 {yaml_file}")
        print(f"錯誤詳情: {str(e)}")
        return None
    
def main():
    """主函數"""
    
    # 定義模板文件列表
    templates = [
        'template/設計變更申請單.docx',
        'template/系統過版申請單.docx',
        'template/系統測試表.docx',
        'template/聯邦銀行廠商進館上線函.docx'
    ]

    # 定義基礎輸出目錄
    base_output_dir = 'output'

    # 檢查必要目錄是否存在
    if not os.path.exists('histData'):
        print("錯誤: histData 目錄不存在")
        return
        
    if not os.path.exists('template'):
        print("錯誤: template 目錄不存在")
        return

    # 獲取所有的申請單 YAML 文件
    application_files = glob.glob('histData/*.yaml')
    
    if not application_files:
        print("警告: 在 histData/ 目錄中找不到任何 *.yaml 檔案")
        return

    # 處理每個申請單
    for application_file in application_files:
        application_data = load_application_data(application_file)
        
        if application_data is None:
            continue
            
        # 從資料中獲取申請單編號和申請單位
        application_number = application_data.get('申請單編號', 'Unknown')
        application_unit = application_data.get('申請單位', 'Unknown')
        print(f"處理申請單: {application_unit}-{application_number}")
        
        for template in templates:
            if os.path.exists(template):
                process_template(template, base_output_dir, application_data)
            else:
                print(f"警告: 模板文件 {template} 不存在")
        
        print()  # 為了輸出美觀，在每個申請單處理完後打印一個空行

if __name__ == '__main__':
    main()
